from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Microsoft YaHei"
st.font.size = Pt(11)

doc.add_heading("CATIA 产品设计报告", level=0)
p = doc.add_paragraph("产品名称：汽车车门转运车架（焊接方管框架）")
p.runs[0].bold = True

# ---------------- 1 产品介绍 ----------------
doc.add_heading("1. 设计产品介绍", level=1)
doc.add_paragraph(
    "本产品是一款用于汽车制造与维修环节的汽车车门转运架（料架）。它是一台由 40×40 焊接方管"
    "构成的框架结构，用于在生产流转、仓储或装配时将多扇汽车车门立式集中存放、转运，避免车门"
    "磕碰并实现快速定位、搬运。框架下方装有带脚轮安装位的底架，整体可移动、可堆叠。"
)
doc.add_paragraph(
    "框架主要由四部分组成：①顶部承托层——11 根沿长度方向均布的纵向梁，形成 10 个车门存放槽；"
    "②中部立柱——13 根方管立柱（4 根角柱顶部伸出作把手），连接上下两层；③底部支撑架——纵梁、"
    "内梁与多道连接短梁构成的底部框架；④定位机构——每个槽底部一根 V 型小梁，车门底边落入 V 槽"
    "实现定位。框架下方另设 8 个“U 型梁+双立柱”门架支撑单元和端部水平支撑管，提高整体刚性与"
    "搬运可靠性。"
)

# ---------------- 2 二维尺寸 ----------------
doc.add_heading("2. 产品二维尺寸", level=1)
doc.add_paragraph("产品主要外形尺寸如下表（单位 mm）。")
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "项目"
hdr[1].text = "尺寸 / 数值 (mm)"
rows = [
    ("总体外形（长×宽×高）", "2400 × 1000 × 800（框架；含把手高 880）"),
    ("方管规格", "40×40（空心，壁厚 2.5）"),
    ("顶部纵向梁", "11 根，40×40×2400，沿宽度方向 20cm 间距均布"),
    ("立柱", "13 根（4 根角柱带把手伸出）"),
    ("V 型定位小梁", "10 根，位于相邻纵梁中间"),
    ("门架支撑单元", "8 个；U 型梁长 580、宽56×高25、厚3；立柱高153"),
    ("端部水平支撑管", "2 根，40×153×250"),
    ("总装配高度", "门架 153 + 框架 960（按题目要求）"),
    ("底部连接短梁", "A/C 间隔各 4、B 间隔 2（40×40 空心方管）"),
]
for a, b in rows:
    c = tbl.add_row().cells
    c[0].text = a
    c[1].text = b

# ---------------- 3 零件截图 ----------------
doc.add_heading("3. CATIA 设计零件截图", level=1)
doc.add_paragraph("主要零件的三维截图如下（各零件均以 .CATPart 文件提交）。")
doc.add_paragraph(
    "框架纵梁（空心方管）—— 40×40×2400，端面为空心矩形；"
    "门架单元（U 型梁 + 双立柱）—— U 型槽 56×25×3，立柱 40×40 空心方管。"
)
doc.add_paragraph(
    "注：各零件的 CATIA 截图可在 CATIA 中逐一打开零件后按“渲染样式-带边着色 / 等轴测”截取，"
    "本报告以二维示意与装配截图配合说明。"
)

# ---------------- 4 总装配截图 ----------------
doc.add_heading("4. CATIA 总装配截图", level=1)
doc.add_paragraph("总装配体三维视图如下（共 78 个组件）。")
doc.add_picture("C:/Users/Administrator/Documents/Codex/2026-08-30/zhao/outputs/assembly_catia.png",
                width=Inches(6.5))

# ---------------- 6 总结 ----------------
doc.add_heading("5. 总结（难点、解决与收获）", level=1)
doc.add_heading("设计难点", level=3)
doc.add_paragraph(
    "① 门架 U 型梁与立柱的装配关系：要求 U 型梁底面朝下贴地、立柱插在 U 型梁内底面、立柱顶面"
    "又要与底部横梁下表面贴合，三者在长度/高度上相互约束，需反复调整；② 杆件之间的干涉：大量"
    "方管在交接处若按中心定位会互相穿透，需按“起点+精确长度+端面贴合”的方式布置；③ 层数多、"
    "零件多（78 个组件），定位与几何关系复杂。"
)
doc.add_heading("解决方案", level=3)
doc.add_paragraph(
    "① 采用子装配/零件分别建模再装配：把立柱做成方形空心管（Pad+Pocket 打通），U 型梁用 U"
    "型截面拉伸，先做出符合要求的门架单元；② 用“按起点+长度”定位坐标，并让连接梁端面贴到"
    "相邻梁内侧面，避免穿透；③ 通过包围盒逐对校验，确保 0 干涉；④ 对装配反复检查并迭代修正"
    "（V 型小梁朝上、门架立柱对齐 AC 短梁等），最终 0 断链、0 干涉。"
)
doc.add_heading("收获与体会", level=3)
doc.add_paragraph(
    "通过本次设计，掌握了 CATIA 参数化建模、方管/型材的拉伸成形（Pad/Pocket 打通做空心件）、"
    "零件与装配体管理、以及装配定位与干涉检查。也体会到“先理解结构、再定尺寸、最后建模”的"
    "重要性；通过不断修正培养了对工程结构布置、公差配合和工艺性的认识。" 
    "零件与装配均以 .CATPart/.CATProduct/.CATDrawing 文件一并提交。"
)

out = "C:/Users/Administrator/Documents/Codex/2026-08-30/zhao/outputs/CATIA_产品设计报告_车门转运架.docx"
doc.save(out)
print("saved", out)
